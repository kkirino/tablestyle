import argparse
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import re
import json
from json.decoder import JSONDecodeError
import sys


def config_args():
    parser = argparse.ArgumentParser(description="utilities for table style in docx")

    # subcommand: get, apply
    subparsers = parser.add_subparsers(title="subcommands")

    # parser for get subcommand
    parser_get = subparsers.add_parser(
        "get", help="get table style parameters and print as standard output"
    )
    parser_get.add_argument("input", type=str, help="path to input docx")
    parser_get.set_defaults(func=get_parameters)

    # parser for apply subcommand
    parser_apply = subparsers.add_parser(
        "apply", help="apply table style parameters according to configuration"
    )
    group_apply = parser_apply.add_mutually_exclusive_group()
    group_apply.add_argument(
        "-n", "--name", type=str, help="name of custom table style to apply"
    )
    group_apply.add_argument("-f", "--file", type=str, help="path to configuration")
    parser_apply.add_argument(
        "-o", "--output", type=str, required=True, help="path to output docx"
    )
    parser_apply.add_argument("input", type=str, help="path to input docx")
    parser_apply.set_defaults(func=apply_parameters)

    return parser.parse_args()


def get_parameters(args):
    try:
        document = Document(args.input)
        for table in document.tables:
            width_list = [col.width for col in table.columns]
            params_dict = {"style": table.style.name, "width": width_list}
            print(json.dumps(params_dict))
        sys.exit(0)
    except FileNotFoundError as e:
        print("FileNotFoundError: {}".format(e), file=sys.stderr)
        sys.exit(1)
    except PackageNotFoundError as e:
        print("PackageNotFoundError: {}".format(e), file=sys.stderr)
        sys.exit(1)


def apply_parameters(args):
    if not re.match(r".+\.docx$", args.output):
        print("ValueError: output file is not docx", file=sys.stderr)
        sys.exit(1)

    try:
        document = Document(args.input)
        if args.file:
            try:
                with open(args.file, mode="r") as reader:
                    for index, line in enumerate(reader):
                        params_dict = json.loads(line)
                        document.tables[index].style = params_dict["style"]
                        for colnum in range(len(params_dict["width"])):
                            table = document.tables[index]
                            column = table.columns[colnum]
                            column.width = params_dict["width"][colnum]
                document.save(args.output)
                sys.exit(0)
            except KeyError as e:
                print("KeyError: {}".format(e), file=sys.stderr)
                sys.exit(1)
            except IndexError:
                print(
                    "IndexError: number of lines in config greater than number of tables",
                    file=sys.stderr,
                )
                sys.exit(1)
        elif args.name:
            try:
                for table in document.tables:
                    table.style = args.name
                document.save(args.output)
                sys.exit(0)
            except KeyError as e:
                print("KeyError: {}".format(e), file=sys.stderr)
                sys.exit(1)
        else:
            try:
                for table in document.tables:
                    table.style = "My Table"
                document.save(args.output)
                sys.exit(0)
            except KeyError as e:
                print("KeyError: {}".format(e), file=sys.stderr)
                sys.exit(1)
    except FileNotFoundError as e:
        print("FileNotFoundError: {}".format(e), file=sys.stderr)
        sys.exit(1)
    except PackageNotFoundError as e:
        print("PackageNotFoundError: {}".format(e), file=sys.stderr)
        sys.exit(1)
    except KeyError as e:
        print(
            "KeyError: configration is not appropreate in key {}".format(e),
            file=sys.stderr,
        )
        sys.exit(1)
    except JSONDecodeError as e:
        print("JSONDecodeError: {}".format(e), file=sys.stderr)
        sys.exit(1)


def main():
    try:
        args = config_args()
        args.func(args)
    except AttributeError:
        print("usage: tablestyle.py [-h] {get,apply} ...", file=sys.stderr)
        print(
            "tablestyle.py: error: the following arguments are required: input",
            file=sys.stderr,
        )
        sys.exit(1)


if __name__ == "__main__":
    main()
